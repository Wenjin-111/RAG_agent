import io
import logging
from langchain_core.documents import Document
from app.ingestion.parsers.base import DocumentParser

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    async def parse(self, data: bytes, file_name: str) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise RuntimeError("python-docx is required for DOCX parsing")

        doc = DocxDocument(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            logger.warning("DOCX parsing produced empty text: %s", file_name)
            return []

        return [Document(
            page_content=full_text.strip(),
            metadata={"source": file_name, "paragraph_count": len(paragraphs)}
        )]
